from docx import Document
from app.utils.logger import get_logger
from datetime import datetime
import os
logger = get_logger(__name__)


class DocxGenerator:

    def generate(self, user, projects, bullets):

        logger.info("Generating Word resume")

        doc = Document()

        # 🧑 Name
        doc.add_heading(user.personal_info.name, 0)

        # 📧 Contact
        doc.add_paragraph(
            f"{user.personal_info.email} | {user.personal_info.phone} | {user.personal_info.location}"
        )

        # 🎓 Education
        doc.add_heading("Education", level=1)
        for edu in user.education:
            doc.add_paragraph(
                f"{edu.degree} - {edu.institution} ({edu.duration})"
            )

        # 💼 Experience
        if user.experience:
            doc.add_heading("Experience", level=1)
            for exp in user.experience:
                doc.add_paragraph(f"{exp.role} - {exp.company}", style="List Bullet")
                for r in exp.responsibilities:
                    doc.add_paragraph(r, style="List Bullet 2")

        # 🚀 Projects
        doc.add_heading("Projects", level=1)
        for proj, bullet in zip(projects, bullets):
            doc.add_paragraph(proj.title, style="List Bullet")
            for b in bullet["bullets"]:
                doc.add_paragraph(b, style="List Bullet 2")

        # 🛠 Skills
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(user.skills.programming_languages))

        now = datetime.now().strftime("%Y%m%d_%H%M%S")
        folder_path = f"generated_resumes/{now}/"
        
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        file_path = f"{folder_path}resume.docx"
        doc.save(file_path)

        logger.info("Word file generated")

        return file_path